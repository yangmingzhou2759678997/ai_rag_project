import os
import unittest

from docx import Document as WordDocument
from docx.table import Table

from tests.test_docx_section_chunking import document_to_bytes, extract_file_text, get_docx_table_text, iter_docx_blocks, split_docx_by_sections
from utils.text_splitter import split_text


class TestDocxTableExtraction(unittest.TestCase):
    def test_paragraph_table_paragraph_order(self):
        document = WordDocument()
        document.add_paragraph("段落A")
        table = document.add_table(rows=1, cols=2)
        table.cell(0, 0).text = "单元格A"
        table.cell(0, 1).text = "单元格B"
        document.add_paragraph("段落B")
        text = extract_file_text("order.docx", document_to_bytes(document))
        self.assertLess(text.index("段落A"), text.index("单元格A"))
        self.assertLess(text.index("单元格B"), text.index("段落B"))

    def test_table_inherits_stage_and_section(self):
        document = WordDocument()
        document.add_paragraph("第三阶段：Agent开发")
        document.add_paragraph("第12天：工具调用")
        table = document.add_table(rows=1, cols=2)
        table.cell(0, 0).text = "工具名称"
        table.cell(0, 1).text = "天气查询"
        chunks = split_docx_by_sections(document_to_bytes(document), 350, 50)
        table_chunks = [item for item in chunks if "天气查询" in item["content"]]
        self.assertEqual(len(table_chunks), 1)
        self.assertIn("第三阶段", table_chunks[0]["stage_title"])
        self.assertIn("第12天", table_chunks[0]["section_title"])

    def test_two_rows_three_columns_and_empty_cell(self):
        document = WordDocument()
        table = document.add_table(rows=2, cols=3)
        values = [["A1", "", "C1"], ["A2", "B2", "C2"]]

        for row_index, row_values in enumerate(values):
            for column_index, value in enumerate(row_values):
                table.cell(row_index, column_index).text = value

        table_text = get_docx_table_text(table)
        self.assertIn("A1 |  | C1", table_text)
        self.assertIn("A2 | B2 | C2", table_text)

    def test_empty_table_returns_empty_text(self):
        document = WordDocument()
        table = document.add_table(rows=2, cols=2)
        self.assertEqual(get_docx_table_text(table), "")

    def test_docx_without_table_is_unchanged(self):
        document = WordDocument()
        document.add_paragraph("第三阶段：基础")
        document.add_paragraph("第12天：内容")
        document.add_paragraph("普通正文")
        file_content = document_to_bytes(document)
        self.assertNotIn("【表格】", extract_file_text("no_table.docx", file_content))
        chunks = split_docx_by_sections(file_content, 350, 50)
        self.assertEqual(chunks[0]["content"], "第三阶段：基础\n第12天：内容\n普通正文")

    def test_table_without_title_uses_original_fallback(self):
        document = WordDocument()
        document.add_paragraph("普通段落A")
        table = document.add_table(rows=1, cols=1)
        table.cell(0, 0).text = "无标题表格内容"
        document.add_paragraph("普通段落B")
        file_content = document_to_bytes(document)
        self.assertIsNone(split_docx_by_sections(file_content, 350, 50))
        chunks = split_text(extract_file_text("plain.docx", file_content), 350, 50)
        self.assertIn("无标题表格内容", "\n".join(chunks))

    def test_preface_and_table_before_first_title_are_kept(self):
        document = WordDocument()
        document.add_paragraph("文档前言")
        table = document.add_table(rows=1, cols=1)
        table.cell(0, 0).text = "前言表格内容"
        document.add_paragraph("第一阶段：基础")
        document.add_paragraph("第1天：开始")
        document.add_paragraph("章节正文")
        chunks = split_docx_by_sections(document_to_bytes(document), 350, 50)
        preface_chunks = [item for item in chunks if "前言表格内容" in item["content"]]
        self.assertEqual(len(preface_chunks), 1)
        self.assertIn("文档前言", preface_chunks[0]["content"])
        self.assertEqual(preface_chunks[0]["stage_title"], "")
        self.assertEqual(preface_chunks[0]["section_title"], "")

    def test_real_docx_contains_table_command_in_original_order(self):
        real_docx_path = os.getenv("TEST_DOCX_PATH", "")

        if not real_docx_path or not os.path.exists(real_docx_path):
            self.skipTest("未设置可用的TEST_DOCX_PATH")

        document = WordDocument(real_docx_path)
        blocks = list(iter_docx_blocks(document))
        table_index = next(index for index, block in enumerate(blocks) if isinstance(block, Table) and get_docx_table_text(block))
        previous_text = next(block.text.strip() for block in reversed(blocks[:table_index]) if not isinstance(block, Table) and block.text.strip())
        table_text = get_docx_table_text(blocks[table_index])
        next_text = next(block.text.strip() for block in blocks[table_index + 1:] if not isinstance(block, Table) and block.text.strip())

        with open(real_docx_path, "rb") as file:
            extracted_text = extract_file_text("real.docx", file.read())

        command_line = next(line for line in table_text.splitlines() if line.startswith("pip ") or line.startswith("git "))
        self.assertIn(command_line, extracted_text)
        self.assertLess(extracted_text.index(previous_text), extracted_text.index(command_line))
        self.assertLess(extracted_text.index(command_line), extracted_text.index(next_text))


if __name__ == "__main__":
    unittest.main()
