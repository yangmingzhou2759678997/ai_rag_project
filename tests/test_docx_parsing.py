import os
import unittest
from io import BytesIO

from docx import Document as WordDocument
from docx.table import Table

from utils.document_parser import extract_document, get_docx_table_text


def document_to_bytes(document):
    file_buffer = BytesIO()
    document.save(file_buffer)
    return file_buffer.getvalue()


class TestDocxParsing(unittest.TestCase):
    def test_paragraph_table_paragraph_order(self):
        document = WordDocument()
        document.add_paragraph("段落A")
        table = document.add_table(rows=1, cols=2)
        table.cell(0, 0).text = "单元格A"
        table.cell(0, 1).text = "单元格B"
        document.add_paragraph("段落B")
        text = extract_document("order.docx", document_to_bytes(document))
        self.assertLess(text.index("段落A"), text.index("单元格A"))
        self.assertLess(text.index("单元格B"), text.index("段落B"))

    def test_two_rows_three_columns_and_empty_cell(self):
        document = WordDocument()
        table = document.add_table(rows=2, cols=3)
        values = [["A1", "", "C1"], ["A2", "B2", "C2"]]

        for row_index, row_values in enumerate(values):
            for column_index, value in enumerate(row_values):
                table.cell(row_index, column_index).text = value

        table_text = get_docx_table_text(table)
        self.assertIn("A1 |  | C1", table_text)
        self.assertIn("A2 | B2 | C2", table_text)

    def test_empty_table_returns_empty_text(self):
        document = WordDocument()
        table = document.add_table(rows=2, cols=2)
        self.assertEqual(get_docx_table_text(table), "")

    def test_headings_are_kept_as_normal_text(self):
        document = WordDocument()
        document.add_paragraph("第一阶段：基础")
        document.add_paragraph("第12天：内容")
        document.add_paragraph("普通正文")
        text = extract_document("headings.docx", document_to_bytes(document))
        self.assertIn("第一阶段：基础", text)
        self.assertIn("第12天：内容", text)
        self.assertIn("普通正文", text)

    def test_high_level_api_returns_paragraph_and_table(self):
        document = WordDocument()
        document.add_paragraph("段落")
        document.add_table(rows=1, cols=1)
        blocks = list(document.iter_inner_content())
        self.assertEqual(len(blocks), 2)
        self.assertFalse(isinstance(blocks[0], Table))
        self.assertTrue(isinstance(blocks[1], Table))

    def test_real_docx_keeps_table_in_original_order(self):
        real_docx_path = os.getenv("TEST_DOCX_PATH", "")

        if not real_docx_path or not os.path.exists(real_docx_path):
            self.skipTest("未设置可用的TEST_DOCX_PATH")

        document = WordDocument(real_docx_path)
        blocks = list(document.iter_inner_content())
        table_index = next(index for index, block in enumerate(blocks) if isinstance(block, Table) and get_docx_table_text(block))
        previous_text = next(block.text.strip() for block in reversed(blocks[:table_index]) if not isinstance(block, Table) and block.text.strip())
        table_text = get_docx_table_text(blocks[table_index])
        next_text = next(block.text.strip() for block in blocks[table_index + 1:] if not isinstance(block, Table) and block.text.strip())

        with open(real_docx_path, "rb") as file:
            extracted_text = extract_document("real.docx", file.read())

        command_line = next(line for line in table_text.splitlines() if line.startswith("pip ") or line.startswith("git "))
        self.assertLess(extracted_text.index(previous_text), extracted_text.index(command_line))
        self.assertLess(extracted_text.index(command_line), extracted_text.index(next_text))


if __name__ == "__main__":
    unittest.main()
