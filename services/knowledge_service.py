# 文件路径: services/knowledge_service.py
import os
import re
from io import BytesIO

from docx import Document as WordDocument
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import Table
from docx.text.paragraph import Paragraph
from pypdf import PdfReader
from sqlalchemy import select
from sqlalchemy.ext.asyncio import AsyncSession

from config import settings
from models import Document
from tools.rag_tool import get_text_embedding
from utils.logger import logger
from utils.text_splitter import split_text


# ==========================================
# DOCX顶层段落和表格提取
# ==========================================
def iter_docx_blocks(word_document):
    """按照DOCX正文XML顺序依次返回顶层段落和表格"""
    for child in word_document.element.body.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, word_document)
        elif isinstance(child, CT_Tbl):
            yield Table(child, word_document)


def get_docx_table_text(table):
    """将表格转换为保留行列顺序的简单文本"""
    row_text_list = []

    for row in table.rows:
        cell_text_list = [cell.text.strip() for cell in row.cells]

        if any(cell_text_list):
            row_text_list.append(" | ".join(cell_text_list))

    if not row_text_list:
        return ""

    return "【表格】\n" + "\n".join(row_text_list)


# ==========================================
# DOCX章节标题识别与切分
# ==========================================
def get_docx_title_type(paragraph):
    """识别阶段标题，以及第X天、章、节标题"""
    text = paragraph.text.strip()
    number_pattern = r"(?:\d+|[零〇一二三四五六七八九十百千两]+)"
    title_match = re.match(rf"^第\s*{number_pattern}\s*(阶段|天|章|节)(?=\s*$|\s*[:：、.\-—]|\s+)", text)

    if not title_match:
        try:
            style_name = paragraph.style.name if paragraph.style else ""
        except KeyError:
            style_name = ""

        if "标题" in style_name or "heading" in style_name.lower():
            title_match = re.match(rf"^第\s*{number_pattern}\s*(阶段|天|章|节)", text)

    if not title_match:
        return ""

    return "stage" if title_match.group(1) == "阶段" else "section"


def split_docx_by_sections(file_content: bytes, chunk_size: int, chunk_overlap: int):
    """按DOCX章节边界切分正文，并让每个片段继承标题"""
    word_document = WordDocument(BytesIO(file_content))
    sections = []
    current_stage_title = ""
    current_section_title = ""
    current_paragraphs = []
    found_title = False

    for block in iter_docx_blocks(word_document):
        text = block.text.strip() if isinstance(block, Paragraph) else get_docx_table_text(block)

        if not text:
            continue

        title_type = get_docx_title_type(block) if isinstance(block, Paragraph) else ""

        if title_type:
            found_title = True

            if current_paragraphs:
                sections.append((current_stage_title, current_section_title, current_paragraphs))
                current_paragraphs = []

            if title_type == "stage":
                current_stage_title = text
                current_section_title = ""
            else:
                current_section_title = text
        else:
            current_paragraphs.append(text)

    if current_paragraphs:
        sections.append((current_stage_title, current_section_title, current_paragraphs))

    if not found_title:
        return None

    section_chunks = []

    for stage_title, section_title, paragraphs in sections:
        body_text = "\n\n".join(paragraphs)
        body_chunks = split_text(text=body_text, chunk_size=chunk_size, chunk_overlap=chunk_overlap)
        title_text = "\n".join(title for title in (stage_title, section_title) if title)

        # chunk_size只控制章节正文大小，标题在切分后附加为检索上下文
        for chunk in body_chunks:
            content = f"{title_text}\n{chunk}" if title_text else chunk
            section_chunks.append({"content": content, "stage_title": stage_title, "section_title": section_title})

    return section_chunks


# ==========================================
# 功能 1：根据文件类型提取文本
# ==========================================
def extract_file_text(file_name: str, file_content: bytes):
    """
    支持提取以下文件中的文字：
    txt、md、pdf、docx
    """
    file_extension = os.path.splitext(file_name)[1].lower()

    # TXT和Markdown都可以直接按UTF-8读取
    if file_extension == ".txt" or file_extension == ".md":
        try:
            return file_content.decode("utf-8-sig")
        except UnicodeDecodeError:
            raise ValueError(
                "TXT或Markdown文件必须使用UTF-8编码"
            )

    # PDF只支持本身带文字层的文件
    if file_extension == ".pdf":
        pdf_reader = PdfReader(BytesIO(file_content))
        page_text_list = []

        for page_index, page in enumerate(pdf_reader.pages):
            page_text = page.extract_text()

            if page_text and page_text.strip():
                page_text_list.append(
                    f"第{page_index + 1}页：\n{page_text.strip()}"
                )

        return "\n\n".join(page_text_list)

    # DOCX按原始顺序提取顶层段落和表格
    if file_extension == ".docx":
        word_document = WordDocument(BytesIO(file_content))
        content_list = []

        for block in iter_docx_blocks(word_document):
            text = block.text.strip() if isinstance(block, Paragraph) else get_docx_table_text(block)

            if text:
                content_list.append(text)

        return "\n\n".join(content_list)

    raise ValueError(
        "当前只支持txt、md、pdf、docx文件"
    )


# ==========================================
# 功能 2：查找同名文件对应的旧知识片段
# ==========================================
async def get_same_file_documents(
    db: AsyncSession,
    file_name: str
):
    """
    查询documents表中的全部片段，
    找出来源文件名相同的片段。
    """
    stmt = select(Document)
    result = await db.execute(stmt)
    all_documents = result.scalars().all()

    same_file_documents = []

    for document in all_documents:
        metadata = document.metadata_info or {}
        source = metadata.get("source", "")

        # seed_db保存的是data/company_knowledge.txt，
        # 上传接口保存的是company_knowledge.txt。
        # 统一只比较最后的文件名。
        source_file_name = os.path.basename(source)

        if source_file_name == file_name:
            same_file_documents.append(document)

    return same_file_documents


# ==========================================
# 功能 3：上传文件并写入向量数据库
# ==========================================
async def save_knowledge_file(
    db: AsyncSession,
    file_name: str,
    file_content: bytes
):
    """
    完成知识文件的文本提取、文本切分、
    向量化和数据库写入。
    """
    # 防止文件名中带有目录路径，只保留最后的文件名
    safe_file_name = os.path.basename(file_name)

    file_extension = os.path.splitext(
        safe_file_name
    )[1].lower().replace(".", "")

    allowed_extensions = settings.allowed_extensions.split(",")
    allowed_extensions = [
        item.strip().lower()
        for item in allowed_extensions
    ]

    # 第一步：校验文件类型
    if file_extension not in allowed_extensions:
        raise ValueError(
            f"不支持.{file_extension}文件，"
            f"只允许上传：{', '.join(allowed_extensions)}"
        )

    # 第二步：校验文件大小
    if len(file_content) > settings.max_upload_size:
        max_size_mb = (
            settings.max_upload_size // 1024 // 1024
        )
        raise ValueError(
            f"文件大小不能超过{max_size_mb}MB"
        )

    # 第三步：提取文件中的文本
    raw_text = extract_file_text(
        safe_file_name,
        file_content
    )

    if not raw_text or not raw_text.strip():
        raise ValueError(
            "没有从文件中提取到有效文本"
        )

    # 第四步：DOCX优先按章节切分，未识别到标题时回退原来的普通切分
    section_chunks = split_docx_by_sections(file_content, settings.chunk_size, settings.chunk_overlap) if file_extension == "docx" else None

    if section_chunks:
        text_chunks = [item["content"] for item in section_chunks]
        chunk_titles = [{"stage_title": item["stage_title"], "section_title": item["section_title"]} for item in section_chunks]
    else:
        text_chunks = split_text(text=raw_text, chunk_size=settings.chunk_size, chunk_overlap=settings.chunk_overlap)
        chunk_titles = [{"stage_title": "", "section_title": ""} for _ in text_chunks]

    if not text_chunks:
        raise ValueError(
            "文件切分后没有生成有效文本片段"
        )

    try:
        # 第五步：如果数据库中已经存在同名文件，
        # 先删除旧片段，避免重复灌入
        old_documents = await get_same_file_documents(
            db,
            safe_file_name
        )

        for old_document in old_documents:
            await db.delete(old_document)

        # 第六步：逐个片段生成向量并写入数据库
        for index, chunk in enumerate(text_chunks):
            logger.info(
                f" [知识库服务] 正在处理文件 "
                f"{safe_file_name} | "
                f"片段 {index + 1}/{len(text_chunks)}"
            )

            embedding_vector = await get_text_embedding(chunk)

            metadata_info = {
                "source": safe_file_name,
                "file_type": file_extension,
                "chunk_index": index
            }

            if file_extension == "docx":
                metadata_info.update(chunk_titles[index])

            new_document = Document(
                content=chunk,
                embedding=embedding_vector,
                metadata_info=metadata_info
            )

            db.add(new_document)

        # 第七步：全部成功后统一提交事务
        await db.commit()

        logger.info(
            f" [知识库服务] 文件 {safe_file_name} "
            f"写入完成，共 {len(text_chunks)} 个片段"
        )

        return {
            "file_name": safe_file_name,
            "file_type": file_extension,
            "chunk_count": len(text_chunks)
        }

    except Exception as e:
        # 中途任何一步失败都回滚，防止数据只写入一部分
        await db.rollback()
        logger.error(
            f" [知识库服务] 文件写入失败：{e}",
            exc_info=True
        )
        raise e


# ==========================================
# 功能 4：查看知识库文件列表
# ==========================================
async def get_knowledge_file_list(db: AsyncSession):
    """
    根据documents表中的metadata_info，
    统计当前知识库包含哪些文件和多少片段。
    """
    stmt = select(Document)
    result = await db.execute(stmt)
    all_documents = result.scalars().all()

    file_dict = {}

    for document in all_documents:
        metadata = document.metadata_info or {}
        source = metadata.get("source", "未知来源")
        file_name = os.path.basename(source)
        file_type = metadata.get("file_type", "txt")

        if file_name not in file_dict:
            file_dict[file_name] = {
                "file_name": file_name,
                "file_type": file_type,
                "chunk_count": 0
            }

        file_dict[file_name]["chunk_count"] += 1

    return list(file_dict.values())


# ==========================================
# 功能 5：删除指定知识库文件
# ==========================================
async def delete_knowledge_file(
    db: AsyncSession,
    file_name: str
):
    """
    根据文件名删除该文件对应的全部知识片段。
    """
    safe_file_name = os.path.basename(file_name)

    try:
        same_file_documents = await get_same_file_documents(
            db,
            safe_file_name
        )

        for document in same_file_documents:
            await db.delete(document)

        await db.commit()

        logger.info(
            f" [知识库服务] 删除文件 "
            f"{safe_file_name}，"
            f"共删除 {len(same_file_documents)} 个片段"
        )

        return len(same_file_documents)

    except Exception as e:
        await db.rollback()
        logger.error(
            f" [知识库服务] 删除文件失败：{e}",
            exc_info=True
        )
        raise e
